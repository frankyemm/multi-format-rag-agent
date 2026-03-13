"""
Document loaders for PDF, DOCX, XLSX, and TXT files.

Each loader extracts text while preserving structure.
XLSX files are converted to Markdown tables via pandas for optimal RAG retrieval.
"""

from pathlib import Path
from typing import Callable

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pandas as pd


# ---------------------------------------------------------------------------
# Individual loaders
# ---------------------------------------------------------------------------

def load_pdf(file_path: str) -> str:
    """Extract text from a PDF preserving page breaks."""
    doc = fitz.open(file_path)
    pages: list[str] = []
    for i, page in enumerate(doc, 1):
        text = page.get_text("text").strip()
        if text:
            pages.append(f"--- Página {i} ---\n{text}")
    doc.close()
    return "\n\n".join(pages)


def load_docx(file_path: str) -> str:
    """Extract text from a DOCX file preserving paragraph structure."""
    doc = DocxDocument(file_path)
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract tables as markdown
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
            paragraphs.append(df.to_markdown(index=False))

    return "\n\n".join(paragraphs)


def load_xlsx(file_path: str) -> str:
    """
    Convert each sheet of an Excel file to a Markdown table via pandas.

    This preserves row-column relationships critical for tabular data
    precision as required by the no-hallucination protocol.
    """
    xls = pd.ExcelFile(file_path)
    sections: list[str] = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        # Drop fully empty rows/columns
        df = df.dropna(how="all").dropna(axis=1, how="all")
        if df.empty:
            continue
        md_table = df.to_markdown(index=False)
        sections.append(f"### Hoja: {sheet_name}\n\n{md_table}")
    return "\n\n".join(sections)


def load_txt(file_path: str) -> str:
    """Read plain text files."""
    return Path(file_path).read_text(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_LOADERS: dict[str, Callable[[str], str]] = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
    ".xls": load_xlsx,
    ".txt": load_txt,
}

SUPPORTED_EXTENSIONS = set(_LOADERS.keys())


def load_document(file_path: str) -> str:
    """
    Dispatch to the correct loader based on file extension.

    Raises ValueError if the file type is not supported.
    """
    ext = Path(file_path).suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(
            f"Tipo de archivo no soportado: '{ext}'. "
            f"Extensiones válidas: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return loader(file_path)
